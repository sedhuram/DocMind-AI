"""One-time fixture generator for parser tests.

Generates:
- backend/tests/core/fixtures/sample.pdf: a real 3-page PDF with extractable
  text on pages 1 and 3, and a blank/near-empty page 2 (to test that
  _parse_pdf preserves original page numbers when skipping empty pages).
- backend/tests/core/fixtures/sample.docx: a real DOCX with 3+ paragraphs,
  including one blank/whitespace-only paragraph (to test that it's dropped).

Requires `reportlab` and `python-docx` installed in backend/.venv.
reportlab is NOT a runtime dependency -- it is only used here to draw text
onto PDF pages for fixture generation, and is intentionally excluded from
requirements.txt / requirements-dev.txt.

Run with: backend/.venv/bin/python generate_fixtures.py
"""

from pathlib import Path

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx

FIXTURES = Path("/Users/sedhuram/Documents/assignment/backend/tests/core/fixtures")
FIXTURES.mkdir(parents=True, exist_ok=True)

# --- sample.pdf: 3 pages, page 2 blank ---
pdf_path = FIXTURES / "sample.pdf"
c = canvas.Canvas(str(pdf_path), pagesize=letter)

# Page 1: real text
c.drawString(72, 720, "DocMind AI Page One")
c.drawString(72, 700, "This document demonstrates PDF parsing with page-aware extraction.")
c.showPage()

# Page 2: intentionally blank (no drawString calls at all)
c.showPage()

# Page 3: real text
c.drawString(72, 720, "DocMind AI Page Three")
c.drawString(72, 700, "This is the third page of the sample PDF fixture for tests.")
c.showPage()

c.save()
print(f"Wrote {pdf_path}")

# --- sample.docx: paragraphs incl. one blank ---
docx_path = FIXTURES / "sample.docx"
document = docx.Document()
document.add_paragraph("DocMind AI supports DOCX ingestion as a first-class format.")
document.add_paragraph("   ")  # blank/whitespace-only paragraph -- must be dropped
document.add_paragraph("Paragraphs are joined together into a single unpaginated page.")
document.save(str(docx_path))
print(f"Wrote {docx_path}")
